#!/usr/bin/env python3
"""Step 28: build a submission-format Word (.docx) from Paper1_Manuscript_FINAL.md,
embedding the 600-dpi figures. Times New Roman 12pt, double-spaced."""
import re, os, sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image

SRC="Paper1_Manuscript_FINAL.md"
_ver = sys.argv[1] if len(sys.argv)>1 else "v2"
OUT=f"Paper1_Manuscript_SUBMISSION_{_ver}.docx"
doc=Document()
# base style: Times New Roman 12, double spaced
st=doc.styles["Normal"]; st.font.name="Times New Roman"; st.font.size=Pt(12)
st.paragraph_format.line_spacing=2.0; st.paragraph_format.space_after=Pt(6)

def add_runs(p, text):
    # inline **bold**, *italic*, `code`
    for piece in re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', text):
        if not piece: continue
        if piece.startswith("**") and piece.endswith("**"):
            r=p.add_run(piece[2:-2]); r.bold=True
        elif piece.startswith("*") and piece.endswith("*"):
            r=p.add_run(piece[1:-1]); r.italic=True
        elif piece.startswith("`") and piece.endswith("`"):
            r=p.add_run(piece[1:-1]); r.font.name="Courier New"
        else:
            p.add_run(piece)

def add_image(path, alt):
    if not os.path.exists(path):
        doc.add_paragraph(f"[missing image: {path}]"); return
    w,h=Image.open(path).size; aspect=w/h
    width=Inches(6.3) if aspect>=1.3 else Inches(3.6)   # wide vs square/portrait
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=width)

lines=open(SRC,encoding="utf-8").read().split("\n")
i=0
while i<len(lines):
    ln=lines[i].rstrip()
    if not ln.strip(): i+=1; continue
    if ln.startswith("---"): i+=1; continue                 # hr
    m=re.match(r'^(#{1,4})\s+(.*)', ln)
    img=re.match(r'^!\[([^\]]*)\]\(([^)]+)\)', ln)
    if img:
        add_image(img.group(2), img.group(1))
    elif m:
        lvl=len(m.group(1)); txt=re.sub(r'\*\*','',m.group(2))
        if lvl==1:
            h=doc.add_heading("",level=0); add_runs(h,txt)
        else:
            h=doc.add_heading(level=min(lvl-1,4)); add_runs(h,txt)
            for r in h.runs: r.font.name="Times New Roman"
    elif ln.startswith("|"):
        # simple table block: collect consecutive | rows
        rows=[]
        while i<len(lines) and lines[i].lstrip().startswith("|"):
            cells=[c.strip() for c in lines[i].strip().strip("|").split("|")]
            if not re.match(r'^[-:\s|]+$', lines[i].strip().strip("|")): rows.append(cells)
            i+=1
        if rows:
            t=doc.add_table(rows=len(rows),cols=len(rows[0])); t.style="Light Grid Accent 1"
            for ri,row in enumerate(rows):
                for ci,cell in enumerate(row):
                    if ci<len(t.rows[ri].cells):
                        cp=t.rows[ri].cells[ci].paragraphs[0]; cp.paragraph_format.line_spacing=1.0
                        add_runs(cp, cell)
                        for rn in cp.runs: rn.font.size=Pt(9)
        continue
    else:
        p=doc.add_paragraph(); add_runs(p, ln)
    i+=1

doc.save(OUT)
print(f"Saved {OUT} ({os.path.getsize(OUT)//1024} KB)")
