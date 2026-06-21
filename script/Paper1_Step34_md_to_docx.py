#!/usr/bin/env python3
"""Step 34: convert a markdown file to a clean submission .docx (no embedded images).
Handles headings, bold/italic/code inline, tables, bullet lists, hr.
Usage: python3 Paper1_Step34_md_to_docx.py INPUT.md OUTPUT.docx
"""
import re, sys
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC, OUT = sys.argv[1], sys.argv[2]
doc=Document()
st=doc.styles["Normal"]; st.font.name="Times New Roman"; st.font.size=Pt(11)
st.paragraph_format.line_spacing=1.5; st.paragraph_format.space_after=Pt(6)

def add_runs(p, text):
    for piece in re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', text):
        if not piece: continue
        if piece.startswith("**") and piece.endswith("**"):
            r=p.add_run(piece[2:-2]); r.bold=True
        elif piece.startswith("*") and piece.endswith("*"):
            r=p.add_run(piece[1:-1]); r.italic=True
        elif piece.startswith("`") and piece.endswith("`"):
            r=p.add_run(piece[1:-1]); r.font.name="Courier New"
        else:
            p.add_run(piece)

lines=open(SRC,encoding="utf-8").read().split("\n")
i=0
while i<len(lines):
    ln=lines[i].rstrip()
    if not ln.strip(): i+=1; continue
    if re.match(r'^-{3,}$', ln): i+=1; continue
    m=re.match(r'^(#{1,4})\s+(.*)', ln)
    if m:
        lvl=len(m.group(1)); txt=re.sub(r'\*\*','',m.group(2))
        h=doc.add_heading(level=min(lvl,4)); add_runs(h,txt)
        for r in h.runs: r.font.name="Times New Roman"
        i+=1; continue
    if ln.lstrip().startswith("|"):
        rows=[]
        while i<len(lines) and lines[i].lstrip().startswith("|"):
            cells=[c.strip() for c in lines[i].strip().strip("|").split("|")]
            if not re.match(r'^[-:\s|]+$', lines[i].strip().strip("|")): rows.append(cells)
            i+=1
        if rows:
            t=doc.add_table(rows=len(rows),cols=len(rows[0])); t.style="Light Grid Accent 1"
            for ri,row in enumerate(rows):
                for ci,cell in enumerate(row):
                    if ci<len(t.rows[ri].cells):
                        cp=t.rows[ri].cells[ci].paragraphs[0]; cp.paragraph_format.line_spacing=1.0
                        add_runs(cp, cell)
                        for rn in cp.runs: rn.font.size=Pt(9)
        continue
    if re.match(r'^\s*[-*]\s+', ln):
        p=doc.add_paragraph(style="List Bullet"); add_runs(p, re.sub(r'^\s*[-*]\s+','',ln))
    else:
        p=doc.add_paragraph(); add_runs(p, ln)
    i+=1

doc.save(OUT)
print(f"Saved {OUT}")
